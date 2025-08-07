import io
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
from dataclasses import dataclass
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ...core.models import User, Goal, JobRecommendation, SkillLevel


@dataclass
class ResumeSection:
    """Individual section of a resume"""
    title: str
    content: List[Dict[str, Any]]
    order: int
    include: bool = True


@dataclass
class ResumeTemplate:
    """Resume template configuration"""
    name: str
    description: str
    sections: List[str]
    style: Dict[str, Any]


class ResumeGenerator:
    """
    Professional resume generator that creates tailored resumes
    based on user profile and target job requirements.
    """
    
    def __init__(self):
        self.templates = self._initialize_templates()
        self.default_sections = [
            "contact", "summary", "experience", "education", 
            "skills", "projects", "certifications"
        ]
    
    def _initialize_templates(self) -> Dict[str, ResumeTemplate]:
        """Initialize predefined resume templates"""
        
        templates = {
            "modern": ResumeTemplate(
                name="Modern Professional",
                description="Clean, ATS-friendly design suitable for most industries",
                sections=["contact", "summary", "experience", "education", "skills", "projects"],
                style={
                    "font_family": "Helvetica",
                    "font_size": 11,
                    "header_color": colors.Color(0.2, 0.3, 0.5),
                    "accent_color": colors.Color(0.3, 0.4, 0.6),
                    "spacing": "normal"
                }
            ),
            "technical": ResumeTemplate(
                name="Technical",
                description="Optimized for software engineering and technical roles",
                sections=["contact", "summary", "skills", "experience", "projects", "education"],
                style={
                    "font_family": "Helvetica",
                    "font_size": 10,
                    "header_color": colors.Color(0.1, 0.2, 0.4),
                    "accent_color": colors.Color(0.2, 0.3, 0.5),
                    "spacing": "compact"
                }
            ),
            "creative": ResumeTemplate(
                name="Creative",
                description="Stylish design for creative and design roles",
                sections=["contact", "summary", "experience", "skills", "projects", "education"],
                style={
                    "font_family": "Helvetica",
                    "font_size": 11,
                    "header_color": colors.Color(0.4, 0.2, 0.3),
                    "accent_color": colors.Color(0.5, 0.3, 0.4),
                    "spacing": "relaxed"
                }
            )
        }
        
        return templates
    
    def generate_resume(self, user: User, target_job: Optional[JobRecommendation] = None,
                       template_name: str = "modern", format_type: str = "pdf") -> io.BytesIO:
        """
        Generate a tailored resume for the user
        """
        
        # Build resume data
        resume_data = self._build_resume_data(user, target_job)
        
        # Get template
        template = self.templates.get(template_name, self.templates["modern"])
        
        # Generate based on format
        if format_type.lower() == "pdf":
            return self._generate_pdf_resume(resume_data, template)
        elif format_type.lower() == "docx":
            return self._generate_docx_resume(resume_data, template)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
    
    def _build_resume_data(self, user: User, target_job: Optional[JobRecommendation] = None) -> Dict[str, Any]:
        """Build structured resume data from user profile"""
        
        # Contact information
        contact_info = {
            "name": user.name,
            "email": user.email,
            "location": user.location or "",
            "phone": "",  # Would need to be added to User model
            "linkedin": "",  # Would need to be added to User model
            "github": "",  # Would need to be added to User model
            "website": ""  # Would need to be added to User model
        }
        
        # Professional summary (tailored to target job if provided)
        summary = self._generate_summary(user, target_job)
        
        # Skills section (prioritized based on target job)
        skills = self._organize_skills(user, target_job)
        
        # Experience section (would be expanded with real data)
        experience = self._format_experience(user)
        
        # Education section
        education = self._format_education(user)
        
        # Projects section (placeholder - would be extracted from user data)
        projects = self._generate_projects_section(user, target_job)
        
        resume_data = {
            "contact": contact_info,
            "summary": summary,
            "skills": skills,
            "experience": experience,
            "education": education,
            "projects": projects,
            "target_job": target_job.title if target_job else None
        }
        
        return resume_data
    
    def _generate_summary(self, user: User, target_job: Optional[JobRecommendation] = None) -> str:
        """Generate professional summary tailored to target role"""
        
        # Base summary components
        experience_years = user.years_experience or 0
        current_role = user.current_role or "Professional"
        
        # Skills emphasis based on target job
        if target_job:
            relevant_skills = []
            user_skills = set(user.skills.keys()) if user.skills else set()
            target_skills = set(target_job.required_skills + target_job.preferred_skills)
            
            # Find overlapping skills
            matching_skills = user_skills.intersection(target_skills)
            relevant_skills = list(matching_skills)[:3]  # Top 3 matching skills
            
            if relevant_skills:
                skills_text = ", ".join(relevant_skills)
                summary = f"{current_role} with {experience_years}+ years of experience specializing in {skills_text}. "
            else:
                summary = f"Experienced {current_role} with {experience_years}+ years in the field. "
        else:
            # Generic summary
            if user.skills:
                top_skills = [skill for skill, level in user.skills.items() 
                             if level in [SkillLevel.ADVANCED, SkillLevel.EXPERT]][:3]
                if top_skills:
                    skills_text = ", ".join(top_skills)
                    summary = f"{current_role} with {experience_years}+ years of experience in {skills_text}. "
                else:
                    summary = f"Experienced {current_role} with {experience_years}+ years of progressive experience. "
            else:
                summary = f"Dedicated {current_role} with {experience_years}+ years of experience. "
        
        # Add value proposition
        summary += "Proven track record of delivering high-quality results and driving business impact. "
        summary += "Strong problem-solving abilities and collaborative team player with excellent communication skills."
        
        return summary
    
    def _organize_skills(self, user: User, target_job: Optional[JobRecommendation] = None) -> Dict[str, List[str]]:
        """Organize skills by category, prioritizing based on target job"""
        
        if not user.skills:
            return {"Technical": [], "Other": []}
        
        # Categorize skills
        technical_keywords = [
            "python", "javascript", "java", "react", "node", "sql", "aws", "docker",
            "git", "html", "css", "typescript", "machine learning", "tensorflow",
            "kubernetes", "angular", "vue", "redux", "graphql", "mongodb"
        ]
        
        technical_skills = []
        soft_skills = []
        other_skills = []
        
        # If target job provided, prioritize matching skills
        if target_job:
            target_skills = set(skill.lower() for skill in 
                              target_job.required_skills + target_job.preferred_skills)
            
            # Prioritize skills that match target job
            prioritized_skills = []
            remaining_skills = []
            
            for skill, level in user.skills.items():
                if skill.lower() in target_skills:
                    prioritized_skills.append(skill)
                else:
                    remaining_skills.append(skill)
            
            all_skills = prioritized_skills + remaining_skills
        else:
            # Sort by skill level (expert -> advanced -> intermediate -> beginner)
            skill_order = {
                SkillLevel.EXPERT: 4,
                SkillLevel.ADVANCED: 3,
                SkillLevel.INTERMEDIATE: 2,
                SkillLevel.BEGINNER: 1
            }
            all_skills = sorted(user.skills.keys(), 
                              key=lambda s: skill_order.get(user.skills[s], 0), 
                              reverse=True)
        
        # Categorize skills
        for skill in all_skills:
            skill_lower = skill.lower()
            if any(keyword in skill_lower for keyword in technical_keywords):
                technical_skills.append(skill)
            elif skill_lower in ["communication", "leadership", "teamwork", "problem solving"]:
                soft_skills.append(skill)
            else:
                other_skills.append(skill)
        
        skills_dict = {}
        if technical_skills:
            skills_dict["Technical"] = technical_skills
        if soft_skills:
            skills_dict["Soft Skills"] = soft_skills
        if other_skills:
            skills_dict["Other"] = other_skills
        
        return skills_dict
    
    def _format_experience(self, user: User) -> List[Dict[str, Any]]:
        """Format work experience (placeholder - would use real user data)"""
        
        # This would typically come from a more detailed user profile
        # For now, create a placeholder based on current role
        
        experience = []
        
        if user.current_role and user.years_experience:
            # Create a sample experience entry
            experience.append({
                "title": user.current_role,
                "company": "Current Company",
                "location": user.location or "Location",
                "start_date": f"{datetime.now().year - user.years_experience}",
                "end_date": "Present",
                "achievements": [
                    "Led development of key features resulting in improved user engagement",
                    "Collaborated with cross-functional teams to deliver projects on time",
                    "Mentored junior team members and contributed to best practices"
                ]
            })
        
        # Add sample previous experience if years > 2
        if user.years_experience and user.years_experience > 2:
            experience.append({
                "title": f"Junior {user.current_role}" if user.current_role else "Previous Role",
                "company": "Previous Company",
                "location": user.location or "Location",
                "start_date": f"{datetime.now().year - user.years_experience}",
                "end_date": f"{datetime.now().year - 2}",
                "achievements": [
                    "Developed foundational skills in core technologies",
                    "Contributed to team projects and learned industry best practices",
                    "Demonstrated strong learning ability and adaptability"
                ]
            })
        
        return experience
    
    def _format_education(self, user: User) -> List[Dict[str, Any]]:
        """Format education section"""
        
        education = []
        
        # Use education from user profile if available
        if user.education:
            for edu in user.education:
                education.append({
                    "degree": edu,
                    "school": "University Name",  # Would need more detailed education model
                    "location": "",
                    "graduation_year": "",
                    "gpa": "",
                    "relevant_coursework": []
                })
        else:
            # Default education entry
            education.append({
                "degree": "Bachelor's Degree",
                "school": "University Name",
                "location": "",
                "graduation_year": f"{datetime.now().year - 2}",
                "gpa": "",
                "relevant_coursework": []
            })
        
        return education
    
    def _generate_projects_section(self, user: User, target_job: Optional[JobRecommendation] = None) -> List[Dict[str, Any]]:
        """Generate projects section (placeholder for demo)"""
        
        projects = []
        
        # Create sample projects based on user skills and target job
        if user.skills:
            # Project 1: Based on top technical skill
            technical_skills = [skill for skill in user.skills.keys() 
                              if any(tech in skill.lower() 
                                   for tech in ["python", "javascript", "react", "java"])]
            
            if technical_skills:
                main_skill = technical_skills[0]
                projects.append({
                    "name": f"{main_skill} Application",
                    "description": f"Full-stack application built with {main_skill} demonstrating proficiency in modern development practices",
                    "technologies": [main_skill] + list(user.skills.keys())[:3],
                    "github_url": "https://github.com/username/project1",
                    "live_url": "https://project1.com",
                    "highlights": [
                        "Implemented responsive design and user authentication",
                        "Integrated with external APIs and databases",
                        "Deployed using modern DevOps practices"
                    ]
                })
        
        # Project 2: Target job specific
        if target_job:
            job_skills = target_job.required_skills[:3]
            projects.append({
                "name": f"{target_job.title.split()[0]} Portfolio Project",
                "description": f"Demonstration project showcasing skills relevant to {target_job.title} role",
                "technologies": job_skills,
                "github_url": "https://github.com/username/project2",
                "live_url": "",
                "highlights": [
                    f"Applied {job_skills[0]} best practices and industry standards",
                    "Focused on performance optimization and scalability",
                    "Documented thoroughly with comprehensive README"
                ]
            })
        
        return projects
    
    def _generate_pdf_resume(self, resume_data: Dict[str, Any], template: ResumeTemplate) -> io.BytesIO:
        """Generate PDF resume using ReportLab"""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, 
                              rightMargin=0.75*inch, leftMargin=0.75*inch,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Build styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=6,
            textColor=template.style["header_color"],
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=12,
            spaceAfter=6,
            textColor=template.style["header_color"],
            borderWidth=1,
            borderColor=template.style["accent_color"],
            borderPadding=3
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=template.style["font_size"],
            spaceAfter=3
        )
        
        story = []
        
        # Name and contact
        contact = resume_data["contact"]
        story.append(Paragraph(contact["name"], title_style))
        
        contact_info = []
        if contact["email"]:
            contact_info.append(contact["email"])
        if contact["location"]:
            contact_info.append(contact["location"])
        if contact["phone"]:
            contact_info.append(contact["phone"])
        
        if contact_info:
            story.append(Paragraph(" | ".join(contact_info), normal_style))
        
        story.append(Spacer(1, 12))
        
        # Professional Summary
        if resume_data["summary"]:
            story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
            story.append(Paragraph(resume_data["summary"], normal_style))
            story.append(Spacer(1, 12))
        
        # Skills
        if resume_data["skills"]:
            story.append(Paragraph("TECHNICAL SKILLS", heading_style))
            for category, skills in resume_data["skills"].items():
                if skills:
                    skills_text = f"<b>{category}:</b> " + ", ".join(skills)
                    story.append(Paragraph(skills_text, normal_style))
            story.append(Spacer(1, 12))
        
        # Experience
        if resume_data["experience"]:
            story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
            for exp in resume_data["experience"]:
                # Job title and company
                job_line = f"<b>{exp['title']}</b> | {exp['company']}"
                if exp.get("location"):
                    job_line += f" | {exp['location']}"
                story.append(Paragraph(job_line, normal_style))
                
                # Dates
                date_line = f"{exp['start_date']} - {exp['end_date']}"
                story.append(Paragraph(date_line, normal_style))
                
                # Achievements
                for achievement in exp.get("achievements", []):
                    story.append(Paragraph(f"• {achievement}", normal_style))
                
                story.append(Spacer(1, 6))
        
        # Projects
        if resume_data["projects"]:
            story.append(Paragraph("PROJECTS", heading_style))
            for project in resume_data["projects"]:
                project_line = f"<b>{project['name']}</b>"
                if project.get("github_url"):
                    project_line += f" | <link href='{project['github_url']}'>GitHub</link>"
                story.append(Paragraph(project_line, normal_style))
                
                story.append(Paragraph(project["description"], normal_style))
                
                if project.get("technologies"):
                    tech_line = f"<b>Technologies:</b> {', '.join(project['technologies'])}"
                    story.append(Paragraph(tech_line, normal_style))
                
                for highlight in project.get("highlights", []):
                    story.append(Paragraph(f"• {highlight}", normal_style))
                
                story.append(Spacer(1, 6))
        
        # Education
        if resume_data["education"]:
            story.append(Paragraph("EDUCATION", heading_style))
            for edu in resume_data["education"]:
                edu_line = f"<b>{edu['degree']}</b> | {edu['school']}"
                if edu.get("graduation_year"):
                    edu_line += f" | {edu['graduation_year']}"
                story.append(Paragraph(edu_line, normal_style))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def _generate_docx_resume(self, resume_data: Dict[str, Any], template: ResumeTemplate) -> io.BytesIO:
        """Generate DOCX resume using python-docx"""
        
        doc = docx.Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Name and contact
        contact = resume_data["contact"]
        
        # Name
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(contact["name"])
        name_run.font.size = docx.shared.Pt(16)
        name_run.bold = True
        
        # Contact info
        contact_info = []
        if contact["email"]:
            contact_info.append(contact["email"])
        if contact["location"]:
            contact_info.append(contact["location"])
        if contact["phone"]:
            contact_info.append(contact["phone"])
        
        if contact_info:
            contact_para = doc.add_paragraph(" | ".join(contact_info))
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Professional Summary
        if resume_data["summary"]:
            doc.add_heading("PROFESSIONAL SUMMARY", level=2)
            doc.add_paragraph(resume_data["summary"])
        
        # Skills
        if resume_data["skills"]:
            doc.add_heading("TECHNICAL SKILLS", level=2)
            for category, skills in resume_data["skills"].items():
                if skills:
                    skills_para = doc.add_paragraph()
                    skills_para.add_run(f"{category}: ").bold = True
                    skills_para.add_run(", ".join(skills))
        
        # Experience
        if resume_data["experience"]:
            doc.add_heading("PROFESSIONAL EXPERIENCE", level=2)
            for exp in resume_data["experience"]:
                # Job title and company
                job_para = doc.add_paragraph()
                job_run = job_para.add_run(f"{exp['title']} | {exp['company']}")
                job_run.bold = True
                if exp.get("location"):
                    job_para.add_run(f" | {exp['location']}")
                
                # Dates
                doc.add_paragraph(f"{exp['start_date']} - {exp['end_date']}")
                
                # Achievements
                for achievement in exp.get("achievements", []):
                    doc.add_paragraph(f"• {achievement}")
        
        # Projects
        if resume_data["projects"]:
            doc.add_heading("PROJECTS", level=2)
            for project in resume_data["projects"]:
                project_para = doc.add_paragraph()
                project_run = project_para.add_run(project['name'])
                project_run.bold = True
                
                doc.add_paragraph(project["description"])
                
                if project.get("technologies"):
                    tech_para = doc.add_paragraph()
                    tech_para.add_run("Technologies: ").bold = True
                    tech_para.add_run(", ".join(project['technologies']))
                
                for highlight in project.get("highlights", []):
                    doc.add_paragraph(f"• {highlight}")
        
        # Education
        if resume_data["education"]:
            doc.add_heading("EDUCATION", level=2)
            for edu in resume_data["education"]:
                edu_para = doc.add_paragraph()
                edu_run = edu_para.add_run(f"{edu['degree']} | {edu['school']}")
                edu_run.bold = True
                if edu.get("graduation_year"):
                    edu_para.add_run(f" | {edu['graduation_year']}")
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def suggest_resume_improvements(self, user: User, target_job: Optional[JobRecommendation] = None) -> List[str]:
        """Suggest improvements to user's resume"""
        
        suggestions = []
        
        # Check for missing contact information
        if not user.email:
            suggestions.append("Add professional email address")
        if not user.location:
            suggestions.append("Include location/city for better job matching")
        
        # Check skills alignment with target job
        if target_job and user.skills:
            user_skills = set(skill.lower() for skill in user.skills.keys())
            required_skills = set(skill.lower() for skill in target_job.required_skills)
            
            missing_required = required_skills - user_skills
            if missing_required:
                suggestions.append(f"Consider developing these required skills: {', '.join(list(missing_required)[:3])}")
        
        # Check experience details
        if not user.years_experience:
            suggestions.append("Add years of experience to better match with job requirements")
        
        # Skills-related suggestions
        if not user.skills:
            suggestions.append("Add technical skills relevant to your target role")
        elif len(user.skills) < 5:
            suggestions.append("Consider adding more relevant technical skills")
        
        # Generic improvement suggestions
        suggestions.extend([
            "Include quantifiable achievements in experience descriptions",
            "Add links to portfolio projects or GitHub profile",
            "Ensure all bullet points start with strong action verbs",
            "Tailor resume summary to highlight most relevant experience",
            "Consider adding relevant certifications or training"
        ])
        
        return suggestions[:7]  # Return top 7 suggestions
    
    def get_resume_templates(self) -> Dict[str, str]:
        """Get available resume templates"""
        return {name: template.description for name, template in self.templates.items()}